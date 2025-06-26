from docx import Document
import os

def read_docx(doc_path):
    """提取 docx 段落與表格內的文字"""
    doc = Document(doc_path)
    text_content = []

    try:
        # 讀取所有段落
        for para in doc.paragraphs:
            text_content.append(para.text)
    except Exception as e:
        print(f"讀取所有段落時發生錯誤: {e}")
    
    try:
        # 讀取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text.strip() for cell in row.cells])  # 用 Tab 分隔欄位
                text_content.append(row_text)  # 將表格內容加入
    except Exception as e:
        print(f"讀取表格內容時發生錯誤: {e}")

    return "\n".join(text_content)  # 轉換成字串返回

def extract_images_from_docx(doc_path):
    """ 從 docx 檔案中提取圖片並儲存 """
    
    # 讀取 docx 文件
    doc = Document(doc_path)

    # 建立存放圖片的資料夾
    if not os.path.exists("extracted_img"):
        os.makedirs("extracted_img")
        output_folder = "extracted_img"
    else:
        output_folder = "extracted_img"

    image_count = 0  # 計數器

    # 讀取圖片並儲存
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:  # 確保是圖片
            image = doc.part.rels[rel].target_part
            image_data = image.blob  # 獲取二進位數據
            ext = image.content_type.split("/")[-1]  # 取得圖片格式 (如 jpg, png)

            # 設定圖片存檔名稱
            image_path = os.path.join(output_folder, f"doc_{image_count}.{ext}")

            # 存檔
            with open(image_path, "wb") as f:
                f.write(image_data)

            print(f"圖片已保存：{image_path}")
            image_count += 1  # 更新計數

    if image_count == 0:
        print("沒有找到任何圖片。")
    else:
        print(f"總共提取 {image_count} 張圖片！")

if __name__ == "__main__":
    doc_path = "sample/sample_doc.docx"
    extracted_text = read_docx(doc_path)
    print(extracted_text)
    extract_images_from_docx(doc_path)
    with open("extracted_txt/docx_output.txt", "w", encoding="utf-8") as txt_file:
        txt_file.write(extracted_text)
